#!/usr/bin/env python3
"""
parsear_distribucion.py — Convierte el Word del docente (tabla RA | Contenidos | Semana)
a JSON estructurado. Detecta texto resaltado (highlight) como marca de "modificado/unificado".

Uso:
    python3 parsear_distribucion.py distribucion.docx [-o distribucion_docente.json]
"""
import argparse, json, re
from docx import Document

RE_UNIDAD = re.compile(r'^Unidad\s+(\d+)\.?\s*(.*)', re.I)
RE_TEMA = re.compile(r'^(\d+\.\d+(?:\.\d+)?)\.?\s*(.*)')

def cell_data(cell):
    items = []
    for p in cell.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        marked = any(r.font.highlight_color for r in p.runs)
        items.append({"texto": txt, "marcado": marked})
    return items

def parse(path):
    doc = Document(path)
    tabla = doc.tables[0]
    seen, semanas = set(), {}
    for row in tabla.rows[1:]:
        sem = row.cells[2].text.strip()
        if not sem or sem in seen:
            continue
        seen.add(sem)
        ras = [i["texto"] for i in cell_data(row.cells[0])]
        unidades = []
        for c in cell_data(row.cells[1]):
            txt = c["texto"]
            m_u, m_t = RE_UNIDAD.match(txt), RE_TEMA.match(txt)
            if m_u:
                unidades.append({"unidad": int(m_u.group(1)), "titulo": m_u.group(2).strip(),
                                 "marcada": c["marcado"], "temas": []})
            elif m_t and unidades:
                unidades[-1]["temas"].append({"codigo": m_t.group(1),
                                              "titulo": m_t.group(2).strip(),
                                              "marcado": c["marcado"]})
            elif unidades and unidades[-1]["temas"]:
                unidades[-1]["temas"][-1]["titulo"] += " — " + txt
            elif unidades:
                unidades[-1].setdefault("notas", []).append(txt)
        semanas[sem] = {"resultados_aprendizaje": ras, "unidades": unidades,
                        "ra_unificados": len(ras) > 1}
    return {"semanas": semanas}

if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('-o', '--out', default='distribucion_docente.json')
    a = ap.parse_args()
    data = parse(a.docx)
    json.dump(data, open(a.out, 'w'), ensure_ascii=False, indent=2)
    total = sum(len(u["temas"]) for s in data["semanas"].values() for u in s["unidades"])
    print(f"✓ {a.out} — {len(data['semanas'])} semanas, {total} temas")
